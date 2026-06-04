from io import BytesIO
import json

from standard_document_assistant.pathing import virtual_to_host_path
from standard_document_assistant.tools import (
    run_format_source_review,
    run_standard_review,
    validate_review_result_schema,
)
from standard_document_assistant.uploads import save_uploaded_file


def test_run_standard_review_from_markdown_writes_artifacts() -> None:
    record = save_uploaded_file(
        original_filename="review.md",
        thread_id="test-review",
        content=(
            "# 测试标准\n\n"
            "## 1 范围\n\n"
            "本文件规定了测试要求。\n\n"
            "## 2 规范性引用文件\n\n"
            "本文件没有规范性引用文件。\n"
        ).encode("utf-8"),
    )

    result = run_standard_review.invoke(
        {
            "content_path": record.virtual_path,
            "output_subdir": "test-review",
            "trace_id": "trace-test-review",
        }
    )

    assert result["status"] == "success"
    assert result["trace_id"] == "trace-test-review"
    assert result["artifacts"]["report"].startswith("/workspace/output/reviews/test-review/")
    assert result["artifacts"]["result"].endswith("_audit_result.json")
    assert result["artifacts"]["trace"].endswith("_audit_trace.json")
    assert result["artifacts"]["manifest"].endswith("_review_manifest.json")

    validation = validate_review_result_schema.invoke({"result_path": result["artifacts"]["result"]})
    assert validation["valid"] is True


def test_run_standard_review_rejects_pdf_without_mineru_markdown() -> None:
    record = save_uploaded_file(
        original_filename="needs-parse.pdf",
        thread_id="test-review",
        content=b"%PDF-1.4\n",
    )

    result = run_standard_review.invoke({"source_path": record.virtual_path})

    assert result["status"] == "failed"
    assert "parse_file_with_mineru" in result["error"]


def test_run_standard_review_flags_missing_normative_references() -> None:
    record = save_uploaded_file(
        original_filename="missing-normative.md",
        thread_id="test-review",
        content=(
            "# 测试标准\n\n"
            "## 1 范围\n\n"
            "本文件规定了测试要求。\n"
        ).encode("utf-8"),
    )

    result = run_standard_review.invoke(
        {
            "content_path": record.virtual_path,
            "output_subdir": "test-review-missing-scope",
        }
    )

    assert result["status"] == "success"
    assert result["summary"]["warn"] >= 1


def test_run_format_source_review_detects_docx_chapter_numbering() -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("1 范围", style="Heading 1")
    doc.add_paragraph("本文件规定了测试要求。")
    doc.add_paragraph("3 术语和定义", style="Heading 1")
    doc.add_paragraph("3.1 测试术语", style="Heading 2")
    buffer = BytesIO()
    doc.save(buffer)

    record = save_uploaded_file(
        original_filename="format-source.docx",
        thread_id="test-review",
        content=buffer.getvalue(),
    )

    result = run_format_source_review.invoke(
        {
            "source_path": record.virtual_path,
            "output_subdir": "test-review-format",
            "trace_id": "trace-test-format",
        }
    )

    assert result["status"] == "success"
    assert result["summary"]["failed"] >= 1
    result_payload = json.loads(
        virtual_to_host_path(result["artifacts"]["result"]).read_text(encoding="utf-8")
    )
    assert any(item["rule_id"] == "DOCX-FMT-001" for item in result_payload["issues"])
    assert result_payload["format_trace"]["evaluations"]
    validation = validate_review_result_schema.invoke({"result_path": result["artifacts"]["result"]})
    assert validation["valid"] is True


def test_run_format_source_review_detects_pdf_chapter_numbering() -> None:
    import fitz

    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "1 范围", fontname="china-s", fontsize=12)
    page.insert_text((72, 96), "本文件规定了测试要求。", fontname="china-s", fontsize=12)
    page.insert_text((72, 120), "3 术语和定义", fontname="china-s", fontsize=12)
    pdf_bytes = pdf.tobytes()
    pdf.close()

    record = save_uploaded_file(
        original_filename="format-source.pdf",
        thread_id="test-review",
        content=pdf_bytes,
    )

    result = run_format_source_review.invoke(
        {
            "source_path": record.virtual_path,
            "output_subdir": "test-review-pdf-format",
            "trace_id": "trace-test-pdf-format",
        }
    )

    assert result["status"] == "success"
    assert result["summary"]["failed"] >= 1
    result_payload = json.loads(
        virtual_to_host_path(result["artifacts"]["result"]).read_text(encoding="utf-8")
    )
    assert any(item["rule_id"] == "DOCX-FMT-001" for item in result_payload["issues"])
    assert result_payload["format_trace"]["source_type"] == "pdf"


def test_review_tools_do_not_expose_runtime_in_schema() -> None:
    assert "runtime" not in run_standard_review.args
    assert "runtime" not in run_format_source_review.args
    assert "runtime" not in validate_review_result_schema.args


def test_quality_gate_respects_zero_max_review_rounds() -> None:
    from standard_document_assistant.graphs.standard_review.nodes.review import quality_gate

    command = quality_gate(
        {
            "review_round": 0,
            "max_review_rounds": 0,
            "insufficient_scopes": ["foreword"],
            "partial_mode": "sectional",
        }
    )

    assert command.goto == "format_review"
